import streamlit as st
from pdf2docx import Converter
from docx import Document
from reportlab.pdfgen import canvas
import tempfile
import os

st.set_page_config(page_title="PDF ⇆ Word Converter", layout="wide")
st.title("Simple PDF to Word / Word to PDF Converter")
st.write("Created by Abhishek")

tab1, tab2 = st.tabs(["PDF --> Word", "Word --> PDF"])

with tab1:
    st.header("Convert PDF ➡️ WORD")
    pdf_file = st.file_uploader("Upload your PDF file", type=["pdf"], key="pdf_to_word")
    if pdf_file:
        st.info(f"File Uploaded: {pdf_file.name} ({pdf_file.size // 1024} KB)")
        if st.button("Convert to WORD"):
            with st.spinner("Converting your PDF to WORD..."):
                progress = st.progress(0)
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                        temp_pdf.write(pdf_file.read())
                        pdf_path = temp_pdf.name

                    docx_path = pdf_path.replace(".pdf", ".docx")
                    
                    cv = Converter(pdf_path)
                    cv.convert(docx_path)
                    cv.close()
                    progress.progress(100)

                    with open(docx_path, "rb") as out_f:
                        st.success("Conversion completed successfully!")
                        st.download_button(label="Download your Word File", data=out_f, file_name="abhishek_converted.docx")
                except Exception as e:
                    st.error(f"Conversion failed: {e}")
                finally:
                    progress.empty()
                    if os.path.exists(pdf_path):
                        os.remove(pdf_path)
                    if os.path.exists(docx_path):
                        os.remove(docx_path)

with tab2:
    st.header("Convert WORD ➡️ PDF")
    word_file = st.file_uploader("Upload your Word file", type=["docx"], key="word_to_pdf")
    if word_file:
        st.info(f"File Uploaded: {word_file.name} ({word_file.size // 1024} KB)")
        if st.button("Convert to PDF"):
            with st.spinner("Converting your Word file to PDF..."):
                progress = st.progress(0)
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                        temp_docx.write(word_file.read())
                        docx_path = temp_docx.name

                    pdf_path = docx_path.replace(".docx", ".pdf")
                    
                    doc = Document(docx_path)
                    c = canvas.Canvas(pdf_path)
                    y = 800

                    for i, p in enumerate(doc.paragraphs):
                        text = p.text
                        if text.strip():
                            c.drawString(50, y, text)
                            y -= 20
                            if y < 50:
                                c.showPage()
                                y = 800
                        progress.progress(int((i + 1) / len(doc.paragraphs) * 100))

                    c.save()

                    with open(pdf_path, "rb") as out_f:
                        st.success("Conversion completed successfully!")
                        st.download_button(label="Download your PDF File", data=out_f, file_name="abhishek_converted.pdf")
                except Exception as e:
                    st.error(f"Conversion failed: {e}")
                finally:
                    progress.empty()
                    if os.path.exists(docx_path):
                        os.remove(docx_path)
                    if os.path.exists(pdf_path):
                        os.remove(pdf_path)

